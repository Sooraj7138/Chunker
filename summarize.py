from unstructured.partition.auto import partition
from langchain_text_splitters import RecursiveCharacterTextSplitter
import ollama
from docx import Document

def extract_pdf_content(pdf_path: str) -> str:
    elements = partition(
        filename=pdf_path,
        strategy="hi_res",
        languages=["eng"]
    )

    content = []

    for el in elements:
        if el.category == "Table":
            table_html = getattr(el.metadata, "text_as_html", None)
            content.append("\n[Table]\n" + (table_html or str(el)))
        else:
            content.append(str(el))

    return "\n\n".join(content)


def chunk_text(text: str):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100
    )
    return splitter.split_text(text)

def summarize_chunk(chunk: str) -> str:
    prompt = f"""
    Summarize the following content clearly and concisely.
    Preserve all the exact contents and table insights.
    Just return only the summary as output, do not reason it like "here is the summary" or anything like that 

    Content:
    {chunk}
    """

    response = ollama.chat(
        model="llama3.2",
        messages=[
            {"role": "system", "content": "You are a professional document summarizer."},
            {"role": "user", "content": prompt}
        ]
    )

    return response["message"]["content"]

def summarize_pdf(pdf_path: str) -> str:
    raw_text = extract_pdf_content(pdf_path)
    chunks = chunk_text(raw_text)

    summaries = []

    for i, chunk in enumerate(chunks):
        print(f"Summarizing chunk {i+1}/{len(chunks)}...")
        summary = summarize_chunk(chunk)
        summaries.append(summary)

    return "\n\n".join(summaries)

def save_summary_to_doc(summary: str, output_path: str):
    doc = Document()
    doc.add_heading("PDF Summary", level=1)

    for paragraph in summary.split("\n\n"):
        doc.add_paragraph(paragraph)

    doc.save(output_path)

if __name__ == "__main__":
    pdf_path = "VTS/G1110-Ed2.1-Use-of-Decision-Support-Tools-for-VTS-Personnel-January-2022.pdf"
    output_doc = "summary.docx"

    final_summary = summarize_pdf(pdf_path)
    save_summary_to_doc(final_summary, output_doc)

    print("✅ Summary saved to", output_doc)
